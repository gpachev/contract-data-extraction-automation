import os
import pytest
from apps.ocr.docx_extractor import extract_docx_text
from docx import Document

def test_extract_docx_text(tmp_path):
    # Create a sample DOCX file
    docx_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Hello World!")
    doc.add_paragraph("Second paragraph.")
    doc.save(docx_path)
    # Extract text
    text = extract_docx_text(str(docx_path))
    assert text is not None
    assert "Hello World!" in text
    assert "Second paragraph." in text
    # Should join paragraphs with newlines
    assert text == "Hello World!\nSecond paragraph."

def test_extract_docx_text_invalid_file(tmp_path):
    # Try to extract from a non-DOCX file
    txt_path = tmp_path / "not_a_docx.txt"
    txt_path.write_text("Not a docx file")
    text = extract_docx_text(str(txt_path))
    assert text is None 